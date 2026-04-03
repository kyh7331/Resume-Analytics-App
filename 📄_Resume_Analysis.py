import streamlit as st
from dotenv import load_dotenv
import os
import anthropic
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor
import io
import json
from pathlib import Path
from supabase import create_client, Client
import requests
from bs4 import BeautifulSoup
from datetime import datetime
import pytz

load_dotenv()
# Auth 상태 초기화
if "user" not in st.session_state:
    st.session_state.user = None
DB_PATH = Path("experiences.json")

supabase: Client = create_client(
    os.getenv("SUPABASE_URL"),
    os.getenv("SUPABASE_KEY")
)

def load_experiences():
    try:
        response = supabase.table("experiences").select("*").order("created_at").execute()
        return response.data
    except Exception as e:
        st.error(f"Error loading experience DB: {e}")
        return []

def load_user_profile():
    try:
        user_id = st.session_state.get("user", {}).get("id")
        if not user_id:
            return {}
        response = supabase.table("user_profiles").select("*").eq("id", user_id).execute()
        if response.data:
            return response.data[0]
        return {}
    except Exception as e:
        return {}

def extract_jd_from_url(url):
    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        response = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(response.text, "html.parser")
        
        # 불필요한 태그 제거
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        
        text = soup.get_text(separator="\n")
        # 빈 줄 정리
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return "\n".join(lines)
    except Exception as e:
        return f"Error: {e}"

def save_experience(experience):
    try:
        supabase.table("experiences").insert(experience).execute()
    except Exception as e:
        st.error(f"Error saving experience: {e}")

def delete_experience(experience_id):
    try:
        supabase.table("experiences").delete().eq("id", experience_id).execute()
    except Exception as e:
        st.error(f"Error deleting experience: {e}")

def load_master_resume():
    try:
        response = supabase.table("master_resume").select("*").order("updated_at", desc=True).limit(1).execute()
        if response.data:
            return response.data[0]["content"]
        return ""
    except Exception as e:
        st.error(f"Error loading master resume: {e}")
        return ""

def save_master_resume(content):
    try:
        # 기존 데이터 삭제 후 새로 저장 (항상 1개만 유지)
        supabase.table("master_resume").delete().neq("id", "00000000-0000-0000-0000-000000000000").execute()
        supabase.table("master_resume").insert({"content": content}).execute()
        return True
    except Exception as e:
        st.error(f"Error saving master resume: {e}")
        return False

def save_application(company, position, jd, analysis, resume):
    try:
        et = pytz.timezone("America/New_York")
        now_et = datetime.now(et)
        response = supabase.table("applications").insert({
            "company_name": company,
            "job_title": position,
            "jd_text": jd,
            "analysis_result": analysis,
            "resume_version": resume,
            "status": "applied",
            "applied_date": now_et.strftime("%Y-%m-%d"),
        }).execute()
        return response.data[0]["id"]
    except Exception as e:
        st.error(f"Error saving application: {e}")
        return None

def load_applications():
    try:
        response = supabase.table("applications").select("*").order("created_at", desc=True).execute()
        return response.data
    except Exception as e:
        st.error(f"Error loading applications: {e}")
        return []

def update_application_status(application_id, status, notes=""):
    try:
        supabase.table("applications").update({
            "status": status,
            "notes": notes
        }).eq("id", application_id).execute()
        supabase.table("application_status_history").insert({
            "application_id": application_id,
            "status": status,
            "notes": notes
        }).execute()
    except Exception as e:
        st.error(f"Error updating status: {e}")

LLM_MODE = os.getenv("LLM_MODE", "local")

if LLM_MODE == "claude":
    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
else:
    from openai import OpenAI
    client = OpenAI(
        base_url="http://localhost:1234/v1",
        api_key="lm-studio"
    )

st.set_page_config(page_title="Resume Analyzer", page_icon="📄", layout="wide")

# Sidebar
lang = st.sidebar.radio("Analysis Language", ["한국어", "English"], horizontal=True)

st.title("📄 Resume Analyzer")
if not st.session_state.get("user"):
    st.warning("Please log in first.")
    st.stop()

# ── Experience Database ──────────────────────────────────────
st.subheader("Experience Database")

if "experiences" not in st.session_state:
    st.session_state.experiences = load_experiences()

if "user_profile" not in st.session_state:
    st.session_state.user_profile = load_user_profile()

if "form_key" not in st.session_state:
    st.session_state.form_key = 0

if st.session_state.get("form_success"):
    st.success("Experience added!")
    st.session_state.form_success = False

with st.form(key=f"experience_form_{st.session_state.form_key}"):
    st.markdown("#### Add Experience")
    col1, col2 = st.columns(2)
    with col1:
        company = st.text_input("Company *")
        department = st.text_input("Department *")
        title = st.text_input("Title / Level *")
    with col2:
        start_date = st.text_input("Start Date (e.g. 2019.03) *")
        end_date = st.text_input("End Date (e.g. 2022.06, or Present)")
        skills = st.text_input("Skills / Competencies (comma separated)")
    project = st.text_input("Project / Role Summary *")
    description = st.text_area("Description *", height=100)

    submitted = st.form_submit_button("Add Experience")
    if submitted:
        if company and department and title and start_date and project and description:
            experience = {
                "company": company,
                "department": department,
                "title": title,
                "start_date": start_date,
                "end_date": end_date,
                "project": project,
                "description": description,
                "skills": skills
            }
            save_experience(experience)
            st.session_state.experiences = load_experiences()
            st.session_state.form_key += 1
            st.session_state.form_success = True
            st.rerun()
        else:
            st.error("Please fill in all required (*) fields.")

if st.session_state.experiences:
    st.markdown("---")
    st.markdown("#### Experience List")
    for i, exp in enumerate(st.session_state.experiences):
        with st.expander(f"{exp['company']} - {exp['department']} ({exp['start_date']} ~ {exp['end_date']})"):
            st.write(f"**Title:** {exp['title']}")
            st.write(f"**Project:** {exp['project']}")
            st.write(f"**Description:** {exp['description']}")
            st.write(f"**Skills:** {exp['skills']}")
            if st.button("Delete", key=f"delete_{i}"):
                delete_experience(exp["id"])
                st.session_state.experiences = load_experiences()
                st.rerun()
else:
    st.info("No experiences added yet. Use the form above to add your experience.")

# ── Resume & JD Input ────────────────────────────────────────
st.markdown("---")
st.subheader("Resume & Job Description")

# 마스터 이력서 로드
if "master_resume" not in st.session_state:
    st.session_state.master_resume = load_master_resume()

# 마스터 이력서 설정 섹션
with st.expander("📌 Master Resume Settings"):
    st.caption("Save your base resume here. It will be automatically used for analysis.")
    master_input = st.text_area(
        "Master Resume",
        value=st.session_state.master_resume,
        height=300,
        key="master_input"
    )
    if st.button("💾 Save Master Resume"):
        if master_input:
            if save_master_resume(master_input):
                st.session_state.master_resume = master_input
                st.success("Master resume saved!")
        else:
            st.error("Please paste your resume first.")

col1, col2 = st.columns(2)

col1, col2 = st.columns(2)
with col1:
    st.markdown("#### Resume")
    resume_file = st.file_uploader("Upload PDF", type=["pdf"])
    resume_text = st.text_area(
        "Or paste text here",
        value=st.session_state.get("master_resume", ""),
        height=300
    )

with col2:
    st.markdown("#### Job Description")
    company_name = st.text_input("Company Name *", placeholder="e.g. PRA Group")
    job_title = st.text_input("Position Title *", placeholder="e.g. Senior Data Analyst - Strategy Analytics")
    jd_default = ""
    if st.session_state.get("use_fetched_jd"):
        jd_default = st.session_state.get("fetched_jd", "")
        st.session_state.use_fetched_jd = False
    elif st.session_state.get("jd_text"):
        jd_default = st.session_state.get("jd_text", "")

    jd_text = st.text_area(
        "Paste JD here",
        value=jd_default,
        height=200
    )
    
    job_url = st.text_input(
        "Or paste Job URL",
        placeholder="Works with company career pages only (LinkedIn/Indeed not supported)"
    )
    
    col_url1, col_url2 = st.columns([2, 1])
    with col_url1:
        fetch_btn = st.button("🔗 Fetch JD from URL", disabled=not job_url)
    with col_url2:
        use_btn = st.button("✅ Use This JD", disabled=not st.session_state.get("fetched_jd"))
    
    if fetch_btn and job_url:
        with st.spinner("Fetching JD..."):
            fetched_jd = extract_jd_from_url(job_url)
            if fetched_jd.startswith("Error"):
                st.error(f"Failed to fetch JD: {fetched_jd}")
            else:
                st.session_state.fetched_jd = fetched_jd
                st.rerun()
    
    if use_btn:
        st.session_state.use_fetched_jd = True
        st.rerun()

if resume_file or resume_text:
    st.session_state.resume_text = resume_text
if jd_text:
    st.session_state.jd_text = jd_text
if company_name:
    st.session_state.company_name = company_name
if job_title:
    st.session_state.job_title = job_title

# ── Functions ────────────────────────────────────────────────
st.markdown("---")

def create_word_report(result, company, position):
    doc = Document()
    title = doc.add_heading("Resume Analysis Report", 0)
    title.alignment = 1
    doc.add_paragraph(f"Company: {company}")
    doc.add_paragraph(f"Position: {position}")
    doc.add_paragraph("")
    for line in result.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=1)
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.replace("**", ""))
            run.bold = True
        elif line.startswith("- "):
            doc.add_paragraph(line.replace("- ", ""), style="List Bullet")
        else:
            doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def extract_pdf_text(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text

with st.expander("🔍 PDF Parse Preview (Dev)"):
    if resume_file and resume_file.size > 0:
        try:
            st.text(extract_pdf_text(resume_file))
        except Exception as e:
            st.error(f"PDF parsing error: {e}")
    else:
        st.info("Upload a PDF to preview parsed text.")

def analyze(resume, jd, company, position, experiences, lang="한국어"):
    experience_text = ""
    for exp in experiences:
        experience_text += f"""
- Company: {exp['company']} / Dept: {exp['department']} / Title: {exp['title']}
  Period: {exp['start_date']} ~ {exp['end_date']}
  Project: {exp['project']}
  Description: {exp['description']}
  Skills: {exp['skills']}
"""

    if lang == "English":
        system_prompt = """You are a realistic career advisor with deep expertise in the US job market.
Your role is to analyze the candidate's resume, experience database, and job description to provide honest, conservative feedback.

Analysis Rules:
1. Never use overly positive expressions. Avoid "excellent", "perfect", "you will definitely hear back".
2. Identify the position title and company type first, then evaluate with appropriate weighting.
   - Senior/Manager level: High weight on leadership, 0→1 experience, strategic thinking
   - Analyst level: High weight on domain knowledge, technical skills, analytical ability
3. Location risk criteria:
   - Use the address on the resume as the base location
   - Over 2 hours drive + not Remote = Relocation required
   - 1-2 hours drive = commute may be difficult
   - Under 2 hours drive = commutable (regardless of Hybrid/Onsite)
   - Remote = no risk
4. Always detect and flag Industry/Function change risks.
5. Find relevant experiences in the experience DB that are not in the resume and recommend them.
6. All resume edit suggestions must be in English. Include Summary Section. Minimum 5 suggestions.
7. Find resume items that are irrelevant or harmful to this specific JD and recommend removal or reduction.
   - For removing entire positions, be conservative. Use "consider removing" language only.
   - For specific bullets or projects, actively recommend removal if irrelevant.
8. When suggesting resume edits, first analyze the existing writing style, vocabulary level, and sentence length. Match that style exactly so edits feel natural and not AI-generated.
9. Classify JD industry proximity as High (same industry) / Functional (similar function) / Distant (different industry) and apply appropriate risk weighting.
10. Prioritize business logic transferability over technical skills (e.g., financial risk modeling → pricing strategy in manufacturing).
11. Evaluate immediate deployability strictly at senior level, but objectively assess learning curve combined with MBA background.
12. No unjustified optimism or praise. Maintain realistic assessment based on data and market logic only.

Results must follow this exact format. Do not change section titles:

## 🏢 Position Analysis
**Company:** [Name]
- Type: [Fortune500/Startup/etc] | Industry: [industry] | HQ: [location]
**Level:** [Junior/Mid/Senior/Manager/Director] | [IC or People Manager]
**Key Requirements:**
- (keyword 1)
- (keyword 2)
- (keyword 3)
**Top Weight:** [keyword 1], [keyword 2]

## ⚠️ Change Risk
#**Industry:** [current] → [target] | Risk: None/Low/Medium/High | [one or two keyword reason]
**Function:** [current] → [target] | Risk: None/Low/Medium/High | [one or two keyword reason]
**Location:** [current] → [target] | [Commutable/Relocation Required/Remote-OK]

## 📊 Resume Fit
**Overall:** (Strong/Moderate/Weak) | (one line summary)
**Strengths:** (one line: what matches JD)
- (keyword 1)
- (keyword 2)
- (keyword 3)
**Critical Gaps:** (one line: what JD requires but resume lacks)
- (gap 1)
- (gap 2)
**Improvement Direction:** (1-2 sentences on how to reframe the resume)
**Remove/Reduce:** (item name + one keyword reason)
- (item): (reason keyword)

## 💡 Experience DB Recommendations
(Experiences not in resume but worth adding from the DB. None if none)
- **[Company - Project]:** [keyword: why relevant]

## ✏️ Resume Edit Suggestions
(Minimum 5 suggestions including Summary Section, all in English)
**1. [What to edit]**
- Before: (first 3-7 words of the existing bullet or summary line to identify it, followed by "...")
- After: (revised English text)
- Why: [keyword]

(repeat for 5+ suggestions)

## 🎯 Final Assessment
**Current pass rate estimate:** (realistic %)
**Post-edit estimate:** (if all suggestions applied)
**Key variable:** (the single most critical factor)
**Final advice:** (whether to apply, where to focus, 2-3 sentences)
"""

    else:
        system_prompt = """당신은 미국 채용 시장에 정통한 현실적인 커리어 어드바이저입니다.
당신의 역할은 구직자가 제공한 이력서, 경력 데이터베이스, Job Description을 분석하여
냉정하고 현실적인 피드백을 제공하는 것입니다.

분석 원칙:
1. 절대 과도하게 긍정적인 표현 금지. "훌륭합니다", "완벽합니다", "연락이 올 것입니다" 사용 금지.
2. 포지션 타이틀과 회사 타입을 먼저 파악하고 그에 맞는 가중치로 평가하세요.
   - Senior/Manager급: 리더십, 0→1 경험, 전략적 사고에 높은 가중치
   - Analyst급: 도메인 지식, 기술 스킬, 분석 능력에 높은 가중치
3. Location 리스크 판단 기준:
   - 이력서의 주소를 기준으로 JD의 근무지까지 운전 거리를 판단하세요
   - 운전 2시간 이상 + Remote가 아닌 경우 = Relocation 필요로 간주
   - 운전 1시간 이상 그러나 2시간 미만인 경우 = 출퇴근이 어려울 수 있음으로 간주
   - 운전 2시간 이내 = 출퇴근 가능으로 간주 (Hybrid/Onsite 무관)
   - Remote = 리스크 없음
4. Industry/Function 변경 리스크를 반드시 감지하고 명시하세요.
5. 이력서에 없지만 경력 DB에 있는 관련 경험을 반드시 발굴하여 추천하세요.
6. 이력서 문구 수정 제안은 반드시 영문으로 작성하세요. Summary Section을 포함하여 최소 5개 이상 제안하세요.
7. 이력서에서 JD와 무관하거나 오히려 방해가 되는 경력/프로젝트를 찾아 제거 또는 축소를 권장하세요.
   - 경력 전체 삭제는 신중하게 권장하세요. "제거 고려" 수준으로만 표현하세요.
   - 특정 프로젝트나 bullet point 축소/제거는 적극적으로 권장하세요.
8. 이력서 수정 문구 제안 시 반드시 기존 이력서의 문체, 어휘 수준, 문장 길이를 먼저 분석하고 동일한 스타일을 유지하세요.
9. JD 산업군을 High (동일 산업) / Functional (유사 기능) / Distant (다른 산업) 으로 분류하고 그에 맞는 리스크 가중치를 적용하세요.
10. 단순 기술 스택보다 비즈니스 로직의 전이 가능성을 우선 평가하세요. (예: 금융 리스크 모델링 → 다른 산업의 가격 전략)
11. 시니어 레벨 기준으로 즉시 투입 가능성은 엄격하게 평가하되, 도메인 학습 곡선을 MBA 경력과 결합해서 현실적으로 산출하세요.
12. 근거 없는 낙관이나 칭찬 배제. 데이터와 시장 논리에 기반한 현실적 평가를 유지하세요.

결과는 반드시 아래 형식을 정확히 따르세요. 섹션 제목을 절대 바꾸지 마세요:

## 🏢 포지션 분석
**회사:** [이름]
- 유형: [대기업/스타트업 등] | 산업: [산업] | 본사: [위치]
**레벨:** [Junior/Mid/Senior/Manager/Director] | [개인기여자 or 팀관리자]
**핵심 요구역량:**
- (키워드 1)
- (키워드 2)
- (키워드 3)
**최우선 평가:** [키워드 1], [키워드 2]

## ⚠️ 변경 리스크
**Industry:** [현재] → [지원] | 리스크: 없음/낮음/중간/높음 | [1~2 단어 이유]
**Function:** [현재] → [지원] | 리스크: 없음/낮음/중간/높음 | [1~2 단어 이유]
**Location:** [현재] → [근무지] | [출퇴근가능/Relocation필요/Remote-OK]

## 📊 이력서 적합도
**종합 평가:** (상/중/하) | (한 줄 요약)
**강점:** (JD와 일치하는 핵심 역량 한 줄 요약)
- (키워드 1)
- (키워드 2)
- (키워드 3)
**치명적 공백:** (JD에서 요구하지만 이력서에 없는 것 한 줄 요약)
- (공백 1)
- (공백 2)
**보완 방향:** (이력서에서 어떤 방향으로 수정하면 좋은지 1-2문장)
**제거/축소 권장:** (항목명 + 이유 한 단어)
- (항목명): (이유 키워드)

## 💡 경력 DB 추천
(이력서에 없지만 경력 DB에서 추가하면 도움될 경력. 없으면 "해당 없음")
- **[회사명 - 프로젝트명]:** 추천 이유 (키워드 중심)

## ✏️ 이력서 수정 제안
(현재 이력서 문구를 이 포지션에 맞게 수정. Summary Section을 포함하여 반드시 5개 이상, 영문으로)
**1. [수정 대상 설명]**
- Before: (어느 부분인지 알 수 있게 첫 3~7단어만 보여주고 "..."으로 생략)
- After: (수정된 영문 문구)
- 이유: (왜 이렇게 바꾸는지 한국어로 단, 키워드 중심)

(5개 이상 반복)

## 🎯 최종 평가
**현재 서류 통과 확률:** (현실적 % 추정)
**수정 후 예상 확률:** (제안 전부 반영 시)
**핵심 변수:** (통과 여부를 결정할 가장 중요한 단 하나의 요소)
**최종 조언:** (지원할지 말지, 어디에 집중할지 2-3문장으로)
"""
    # 프로필 컨텍스트 생성
    profile = st.session_state.get("user_profile", {})
    profile_context = ""
    if profile:
        profile_context = f"""
**Candidate Profile:**
- Location: {profile.get('current_location', 'Unknown')}
- Willing to Relocate: {profile.get('willing_to_relocate', 'Unknown')}
- Relocation Cities: {profile.get('relocation_cities', 'None specified')}
- Sponsorship Required: {profile.get('require_sponsorship', 'Unknown')}
- Security Clearance: {profile.get('security_clearance', 'None')}
- Target Salary: ${profile.get('salary_min', 0):,} - ${profile.get('salary_max', 0):,}
- Education: {profile.get('degree', '')} in {profile.get('major', '')} from {profile.get('school', '')} ({profile.get('graduation_year', '')})
- Target Industries: {profile.get('target_industries', 'Not specified')}
- Target Functions: {profile.get('target_functions', 'Not specified')}
- Target Level: {profile.get('target_title_level', 'Not specified')}
"""
    user_message = f"""Please analyze the following:

**Target Company:** {company}
**Target Position:** {position}
{profile_context}
**Current Resume:**
{resume}

**Experience Database (not in resume):**
{experience_text if experience_text else "No additional experiences entered"}

*Job Description:**
{jd}""" if lang == "English" else f"""다음 정보를 바탕으로 분석해주세요:

**지원 회사:** {company}
**지원 포지션:** {position}
{profile_context}
**현재 이력서:**
{resume}

**경력 데이터베이스 (이력서 외 추가 경력):**
{experience_text if experience_text else "입력된 추가 경력 없음"}

**Job Description:**
{jd}"""

    if LLM_MODE == "claude":
        with client.messages.stream(
            model="claude-sonnet-4-5",
            max_tokens=3000,
            system=system_prompt,
            messages=[{"role": "user", "content": user_message}]
        ) as stream:
            for text in stream.text_stream:
                yield text
    else:
        stream = client.chat.completions.create(
            model="gemma-4-26b",
            max_tokens=3000,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ],
            stream=True
        )
        for chunk in stream:
            if chunk.choices[0].delta.content is not None:
                yield chunk.choices[0].delta.content

# ── Analyze Button ───────────────────────────────────────────
if st.button("🔍 Analyze", type="primary"):
    if not resume_text and not resume_file:
        st.error("Please upload or paste your resume.")
    elif not jd_text:
        st.error("Please paste the job description.")
    elif not company_name:
        st.error("Please enter the company name.")
    elif not job_title:
        st.error("Please enter the position title.")
    else:
        if resume_file and resume_file.size > 0:
            resume_file.seek(0)
            final_resume = extract_pdf_text(resume_file)
        elif resume_text:
            final_resume = resume_text
        else:
            final_resume = st.session_state.get("master_resume", "")

        st.session_state.pop("result", None)
        st.session_state["final_resume"] = final_resume
        st.session_state["jd_text"] = jd_text
        st.session_state["company_name"] = company_name
        st.session_state["job_title"] = job_title
        st.session_state["run_analysis"] = True
        st.rerun()

if st.session_state.get("run_analysis"):
    st.session_state["run_analysis"] = False
    result_placeholder = st.empty()
    full_result = ""

    for chunk in analyze(
        resume=st.session_state["final_resume"],
        jd=st.session_state["jd_text"],
        company=st.session_state["company_name"],
        position=st.session_state["job_title"],
        experiences=st.session_state.experiences,
        lang=lang
    ):
        full_result += chunk
        result_placeholder.markdown(full_result + "▌")

    result_placeholder.empty()
    st.session_state.result = full_result
    st.rerun()

def format_result(text):
    replacements = [
        # 포지션 분석 - 한글
        ("레벨:", "\n\n**레벨:**"),
        ("핵심 요구역량:", "\n\n**핵심 요구역량:**"),
        ("최우선 평가:", "\n\n**최우선 평가:**"),
        # 포지션 분석 - 영문
        ("Level:", "\n\n**Level:**"),
        ("Key Requirements:", "\n\n**Key Requirements:**"),
        ("Top Weight:", "\n\n**Top Weight:**"),
        # 리스크 섹션 - 한글
        ("Function:", "\n\n**Function:**"),
        ("Location:", "\n\n**Location:**"),
        # 파이프 구분자를 줄바꿈으로
        (" | 레벨:", "\n\n**레벨:**"),
        (" | 핵심", "\n\n**핵심"),
        (" | Level:", "\n\n**Level:**"),
        (" | Key", "\n\n**Key"),
        (" | Top", "\n\n**Top"),
        # 이력서 적합도 - 영문
        ("**Strengths:**", "\n\n**Strengths:**"),
        ("**Critical Gaps:**", "\n\n**Critical Gaps:**"),
        ("**Improvement Direction:**", "\n\n**Improvement Direction:**"),
        ("**Remove/Reduce:**", "\n\n**Remove/Reduce:**"),
        # 최종 평가 - 한글
        ("**현재 서류 통과 확률:**", "\n\n**현재 서류 통과 확률:**"),
        ("**수정 후 예상 확률:**", "\n\n**수정 후 예상 확률:**"),
        ("**핵심 변수:**", "\n\n**핵심 변수:**"),
        ("**최종 조언:**", "\n\n**최종 조언:**"),
        # 최종 평가 - 영문
        ("**Current pass rate estimate:**", "\n\n**Current pass rate estimate:**"),
        ("**Post-edit estimate:**", "\n\n**Post-edit estimate:**"),
        ("**Key variable:**", "\n\n**Key variable:**"),
        ("**Final advice:**", "\n\n**Final advice:**"),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    return text

# ── Results ──────────────────────────────────────────────────
if "result" in st.session_state:
    st.components.v1.html("""
    <script>
    function copyToClipboard(text) {
        navigator.clipboard.writeText(text).then(function() {
            alert('Copied!');
        });
    }
    </script>
    """, height=0)

    st.markdown("---")
    st.subheader("📋 Analysis Results")

    result = st.session_state.result
    result = format_result(st.session_state.result)

    if lang == "English":
        sections = {
            "🏢 Position Analysis": "",
            "⚠️ Change Risk": "",
            "📊 Resume Fit": "",
            "💡 Experience DB Recommendations": "",
            "✏️ Resume Edit Suggestions": "",
            "🎯 Final Assessment": ""
        }
    else:
        sections = {
            "🏢 포지션 분석": "",
            "⚠️ 변경 리스크": "",
            "📊 이력서 적합도": "",
            "💡 경력 DB 추천": "",
            "✏️ 이력서 수정 제안": "",
            "🎯 최종 평가": ""
        }

    current_section = None
    for line in result.split("\n"):
        for section in sections.keys():
            if section in line:
                current_section = section
                break
        if current_section:
            sections[current_section] += line + "\n"
    
    with st.expander("🔍 섹션 파싱 디버그"):
        for k, v in sections.items():
            st.write(f"**{k}:** {len(v)} chars")
            st.write(v[:200] if v else "EMPTY")

    # Download
    col1, col2 = st.columns([1, 4])
    with col1:
        filename = f"{st.session_state.get('company_name', 'company')}_{st.session_state.get('job_title', 'position')}.docx"
        filename = filename.replace(" ", "_")
        word_file = create_word_report(
            result=result,
            company=st.session_state.get('company_name', ''),
            position=st.session_state.get('job_title', '')
        )
        st.download_button(
            label="📥 Download as Word",
            data=word_file,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # Application tracking
    st.markdown("---")
    col1, col2 = st.columns([1, 4])
    with col1:
        if st.button("✅ I Applied to This Position", type="primary"):
            app_id = save_application(
                company=st.session_state.get("company_name", ""),
                position=st.session_state.get("job_title", ""),
                jd=st.session_state.get("jd_text", ""),
                analysis=format_result(st.session_state.result),
                resume=st.session_state.get("resume_text", "")
            )
            if app_id:
                st.session_state.last_application_id = app_id
                st.success("Application saved!")

    if "last_application_id" in st.session_state:
        st.markdown("#### 📋 Update Application Status")
        status_options = {
            "applied": "📝 Applied",
            "screening": "📞 Screening Call",
            "interview": "🤝 Interview",
            "offer": "🎉 Offer",
            "rejected": "❌ Rejected",
            "withdrawn": "🚫 Withdrawn"
        }
        col1, col2 = st.columns(2)
        with col1:
            new_status = st.selectbox(
                "Current Status",
                options=list(status_options.keys()),
                format_func=lambda x: status_options[x]
            )
        with col2:
            status_notes = st.text_input("Notes (optional)", placeholder="e.g. Passed 1st interview, scheduling 2nd round")

        if st.button("Update Status"):
            update_application_status(
                st.session_state.last_application_id,
                new_status,
                status_notes
            )
            st.success("Status updated!")

    # View mode
    view_mode = st.radio("View Mode", ["Tab View", "Full View"], horizontal=True)

    if view_mode == "Full View":
        st.markdown(result)
    else:
        if lang == "English":
            tab_labels = ["🏢 Position", "⚠️ Risk", "📊 Fit", "💡 DB Recs", "✏️ Edits", "🎯 Final"]
        else:
            tab_labels = ["🏢 포지션 분석", "⚠️ 변경 리스크", "📊 적합도", "💡 경력 추천", "✏️ 문구 수정", "🎯 최종 평가"]

        tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(tab_labels)
        keys = list(sections.keys())

        with tab1:
            st.markdown(sections[keys[0]])
        with tab2:
            content = sections[keys[1]]
            if content.strip():
                st.markdown(content)
            else:
                st.success("No change risk detected" if lang == "English" else "변경 리스크 없음")
        with tab3:
            st.markdown(sections[keys[2]])
        with tab4:
            st.markdown(sections[keys[3]])
        with tab5:
            content = sections[keys[4]]
            st.markdown(content)
            
            # Before/After 파싱해서 복사 버튼 추가
            lines = content.split("\n")
            after_texts = []
            for line in lines:
                if line.strip().startswith("- After:"):
                    after_text = line.replace("- After:", "").strip()
                    after_texts.append(after_text)
            
            if after_texts:
                st.markdown("---")
                st.markdown("#### 📋 Copy After Suggestions")
                for i, text in enumerate(after_texts):
                    col1, col2 = st.columns([4, 1])
                    with col1:
                        st.text(text[:100] + "..." if len(text) > 100 else text)
                    with col2:
                        escaped = text.replace('"', '\\"').replace("'", "\\'").replace("\n", "\\n")
                        st.components.v1.html(f"""
                        <button onclick="navigator.clipboard.writeText('{escaped}').then(()=>this.innerText='✅').catch(()=>this.innerText='❌')" 
                        style="background:#4CAF50;color:white;border:none;padding:5px 10px;border-radius:5px;cursor:pointer;font-size:12px">
                        Copy
                        </button>
                        """, height=35)
        with tab6:
            st.markdown(sections[keys[5]])